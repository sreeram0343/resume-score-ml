import io
from docx import Document
from .base_parser import BaseParser, ParseError
from .schemas import ParseResult, ATSFlags

class DOCXParser(BaseParser):
    """Parser for DOCX files using python-docx."""

    def parse(self, file_bytes: bytes) -> ParseResult:
        self.validate_size(file_bytes)
        
        try:
            doc = Document(io.BytesIO(file_bytes))
            text_parts = []
            ats_flags = ATSFlags()

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables
            if doc.tables:
                ats_flags.tables_detected = True
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            text_parts.append(" | ".join(row_text))

            # Image detection
            # InlineShapes can be checked for pictures
            for shape in doc.inline_shapes:
                # 3 is wdInlineShapePicture, 4 is wdInlineShapeLinkedPicture
                if shape.type in (3, 4):
                    ats_flags.images_detected = True
                    break

            full_text = "\n".join(text_parts)
            ats_flags.special_chars_ratio = self.calculate_special_chars_ratio(full_text)

            return ParseResult(
                text=full_text,
                word_count=len(full_text.split()),
                page_count=1,  # DOCX doesn't have a reliable page count without rendering
                ats_flags=ats_flags,
                parser_used="python-docx",
                warnings=["Page count for DOCX is an estimate."]
            )
            
        except Exception as e:
            raise ParseError(f"DOCX parsing failed: {str(e)}", file_type="DOCX")
