import pytest
import io
from docx import Document
from app.parsers.dispatcher import ParserDispatcher
from app.parsers.base_parser import ParseError

def test_valid_txt_extraction():
    content = b"This is a simple resume text."
    result = ParserDispatcher.parse(content)
    assert result.parser_used == "txt"
    assert "simple resume" in result.text
    assert result.word_count == 6

def test_valid_docx_extraction():
    # Create a simple DOCX in memory
    doc_io = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Experience at Google")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Skill"
    table.cell(0, 1).text = "Python"
    doc.save(doc_io)
    
    result = ParserDispatcher.parse(doc_io.getvalue())
    assert result.parser_used == "python-docx"
    assert "Experience at Google" in result.text
    assert "Skill | Python" in result.text
    assert result.ats_flags.tables_detected is True

def test_empty_file():
    with pytest.raises(ParseError) as excinfo:
        ParserDispatcher.parse(b"")
    assert "Empty file" in str(excinfo.value)

def test_file_type_spoofing():
    # TXT content with .pdf extension (though dispatcher uses bytes, not extension)
    content = b"Not a PDF really"
    result = ParserDispatcher.parse(content)
    # Should fallback to TXT since magic bytes don't match PDF
    assert result.parser_used == "txt"

@pytest.mark.skip(reason="Requires valid PDF binary for testing")
def test_pdf_extraction():
    pass
